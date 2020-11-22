from PIL import Image
import os,sys
from docx import Document
from docx.shared import Cm
import math

class QCMDisplay():
    def __init__(self):
        if len(sys.argv) >= 3:
            ext = sys.argv[1]
            column = int(sys.argv[2])
            if len(sys.argv)==4:
                self.file_type = sys.argv[3]
            else:
                self.file_type = "tex"
        else:
            print("Error : not enough argument")
            sys.exit(1)
        liste = os.listdir(".")
        self.files = []
        for _ in liste:
            if ext in _:
                self.files.append(_)
        self.files.sort()
        self.column = column
        self.content = '''\\documentclass{article}
\\usepackage[utf8]{inputenc}
\\usepackage[T1]{fontenc}
\\usepackage[margin=2cm]{geometry}
\\usepackage{graphicx}
\\begin{document}'''

    def execute(self):
        self.set_layout()
        self.generate_doc()
        self.compile()

    def set_layout(self):
        self.count_files = len(self.files)
        # margin = 2cm
        # available size roughly 16
        self.width = 16/self.column
        
    def get_size(self):
        list_of_sizes = [[],[]]
        for _ in self.files:
            im = open(_)
            list_of_sizes[0].append(im.size[0])
            list_of_sizes[1].append(im.size[1])
            im.close()
            
    def generate_doc(self):
        count = 0
        for _ in self.files:
            self.content += '''\\noindent\\begin{minipage}{%scm}
            \\includegraphics[width=%scm]{%s}
\\end{minipage}'''%(self.width,self.width,_)
            count += 1
            if count % self.column == 0:
                count = 0
                self.content += '''\\newline'''

        self.content += '''\\end{document}'''


    def make_docx(self):
        document = Document()
        #changing the page margins
        sections = document.sections
        margin = 2
        for section in sections:
            section.top_margin = Cm(margin)
            section.bottom_margin = Cm(margin)
            section.left_margin = Cm(margin)
            section.right_margin = Cm(margin)
        document.add_heading('Document Title', 0)

        p = document.add_paragraph('A plain paragraph having some ')

        document.add_paragraph('Intense quote', style='Intense Quote')


        rows = math.ceil(len(self.files)/self.column)

        paragraph = document.add_paragraph()
        count = 0
        for _ in self.files:
            run = paragraph.add_run()
            run.add_picture(_, width=Cm(self.width))
            count += 1
            if count == self.column:
                count = 0
                paragraph = document.add_paragraph()

        document.save('demo.docx')


        
    def compile(self):
        if self.file_type == "tex":
            content_file = open("content.tex","w")
            content_file.write(self.content)
            content_file.close()
            os.system("pdflatex content.tex")
            os.system("rm content.aux content.log")
        else:
            self.make_docx()

if __name__=="__main__":
    app = QCMDisplay()
    app.execute()
